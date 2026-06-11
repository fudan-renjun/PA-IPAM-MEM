#!/usr/bin/env python3
"""Create an NC-style supplementary material Word file with embedded tables."""

from __future__ import annotations

import importlib.util
from pathlib import Path

import pandas as pd
from docx import Document
from docx.enum.section import WD_ORIENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt


PROJECT = Path(__file__).resolve().parents[1]
DATE = "2026-06-09"
TABLE_DIR = next(PROJECT.glob("*/manuscript_tables/supplementary_tables"))
OUT_DIR = next(PROJECT.glob("*/manuscript_drafts"))
WRITER_PATH = PROJECT / "scripts" / "18_write_methods_results_manuscript_docs.py"
OUT_PATH = OUT_DIR / f"carbapenem_predictability_supplementary_material_NC_style_{DATE}.docx"


def load_writer():
    spec = importlib.util.spec_from_file_location("writer18", WRITER_PATH)
    module = importlib.util.module_from_spec(spec)
    assert spec.loader is not None
    spec.loader.exec_module(module)
    return module


def make_landscape(doc: Document) -> None:
    for section in doc.sections:
        section.orientation = WD_ORIENT.LANDSCAPE
        section.page_width, section.page_height = section.page_height, section.page_width
        section.top_margin = Pt(36)
        section.bottom_margin = Pt(36)
        section.left_margin = Pt(36)
        section.right_margin = Pt(36)


def shade_cell(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text: str, *, bold: bool = False, font_size: float = 6.2) -> None:
    cell.text = "" if text is None else str(text)
    for paragraph in cell.paragraphs:
        paragraph.paragraph_format.space_after = Pt(0)
        for run in paragraph.runs:
            run.bold = bold
            run.font.name = "Arial"
            run.font.size = Pt(font_size)


def fmt(value) -> str:
    if pd.isna(value):
        return ""
    if isinstance(value, float):
        if abs(value - round(value)) < 1e-9:
            return str(int(round(value)))
        return f"{value:.3g}"
    return str(value)


def add_note(doc: Document, text: str) -> None:
    p = doc.add_paragraph(text)
    p.paragraph_format.space_after = Pt(6)
    for run in p.runs:
        run.font.name = "Arial"
        run.font.size = Pt(8)


def add_table_caption(doc: Document, label: str, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(label)
    r.bold = True
    r.font.name = "Arial"
    r.font.size = Pt(9)
    r = p.add_run(" " + text)
    r.font.name = "Arial"
    r.font.size = Pt(9)


def add_dataframe(doc: Document, df: pd.DataFrame, *, font_size: float = 5.8) -> None:
    clean = df.copy()
    clean = clean.replace({pd.NA: "", None: ""})
    table = doc.add_table(rows=1, cols=len(clean.columns))
    table.style = "Table Grid"
    table.autofit = True
    for j, col in enumerate(clean.columns):
        shade_cell(table.rows[0].cells[j], "D9EAF7")
        set_cell_text(table.rows[0].cells[j], str(col), bold=True, font_size=font_size)
    for _, row in clean.iterrows():
        cells = table.add_row().cells
        for j, col in enumerate(clean.columns):
            set_cell_text(cells[j], fmt(row[col]), font_size=font_size)
    doc.add_paragraph()


def read_table(stem: str, sheet: str) -> pd.DataFrame:
    path = next(TABLE_DIR.glob(f"{stem}*.xlsx"))
    return pd.read_excel(path, sheet_name=sheet)


def add_sheet(doc: Document, label: str, text: str, stem: str, sheet: str, *, font_size: float = 5.8) -> None:
    add_table_caption(doc, label, text)
    add_dataframe(doc, read_table(stem, sheet), font_size=font_size)


def selected(df: pd.DataFrame, columns: list[str]) -> pd.DataFrame:
    return df[[c for c in columns if c in df.columns]].copy()


def add_large_table_compact_views(doc: Document) -> None:
    add_table_caption(
        doc,
        "Supplementary Table S5A.",
        "Compact high-confidence mechanism evidence counts by subtype. The full isolate-level table is provided in the corresponding Excel file.",
    )
    hce = read_table("Supplementary_Table_S05_", "high_confidence_evidence")
    count_cols = [
        "oprd_severe_loss",
        "oprd_deep_disruptive",
        "efflux_strict_driver_disruptive_any",
        "ampc_core_driver_disruptive_any",
        "acquired_carbapenemase_strict",
    ]
    for col in count_cols:
        if col not in hce.columns:
            hce[col] = 0
    hce_summary = (
        hce.groupby("mechanism_subtype_high_confidence", dropna=False)
        .agg(
            isolates=("genome_id", "nunique"),
            oprd_severe_loss=("oprd_severe_loss", "sum"),
            oprd_deep_disruptive=("oprd_deep_disruptive", "sum"),
            efflux_driver_disruptive=("efflux_strict_driver_disruptive_any", "sum"),
            ampc_axis_disruptive=("ampc_core_driver_disruptive_any", "sum"),
            acquired_carbapenemase=("acquired_carbapenemase_strict", "sum"),
        )
        .reset_index()
    )
    add_dataframe(doc, hce_summary, font_size=6.2)

    add_table_caption(
        doc,
        "Supplementary Table S5B.",
        "Compact prediction-error summary by high-confidence mechanism subtype. Full isolate-level prediction rows are provided in Excel.",
    )
    pred = read_table("Supplementary_Table_S05_", "prediction_with_mechanism")
    pred_summary = (
        pred.groupby(["drug", "mechanism_subtype_high_confidence"], dropna=False)
        .agg(
            n=("genome_id", "size"),
            within_1_dilution=("mic_ea_pm1", "mean"),
            mae_log2=("abs_log2_error", "mean"),
            bias_log2=("signed_log2_error", "mean"),
            categorical_agreement=("sns_correct", "mean"),
            false_susceptible_n=("vme", "sum"),
            false_resistant_n=("me", "sum"),
        )
        .reset_index()
    )
    for col in ["within_1_dilution", "categorical_agreement"]:
        pred_summary[col] = pred_summary[col] * 100
    add_dataframe(doc, pred_summary, font_size=6.0)

    add_table_caption(
        doc,
        "Supplementary Table S6A.",
        "Target-gene variation burden summary. Full target-level and variant-row data are provided in Excel.",
    )
    tgt = read_table("Supplementary_Table_S06_", "target_gene_variant_summary")
    burden = (
        tgt.groupby(["axis", "gene", "locus_tag", "strict_driver"], dropna=False)
        .agg(
            isolates=("genome_id", "nunique"),
            isolates_with_variant=("variant_count", lambda s: int((s > 0).sum())),
            isolates_with_non_syn_or_regulatory=("n_non_syn_or_regulatory", lambda s: int((s > 0).sum())),
            isolates_with_disruptive=("n_disruptive", lambda s: int((s > 0).sum())),
            total_variants=("variant_count", "sum"),
        )
        .reset_index()
        .sort_values(["axis", "gene"])
    )
    add_dataframe(doc, burden, font_size=5.8)

    add_table_caption(
        doc,
        "Supplementary Table S6B.",
        "Most frequent target-gene variant signatures, shown as a compact preview of the full variant-row table.",
    )
    variant_rows = read_table("Supplementary_Table_S06_", "target_variant_rows")
    top_variants = (
        variant_rows.groupby(["gene", "location_class", "variant_type", "effect_class", "effect"], dropna=False)
        .size()
        .reset_index(name="rows")
        .sort_values("rows", ascending=False)
        .head(30)
    )
    add_dataframe(doc, top_variants, font_size=5.6)


def add_supplementary_tables(doc: Document) -> None:
    doc.add_page_break()
    doc.add_heading("Supplementary Tables", level=1)
    add_note(
        doc,
        "Tables are formatted for review readability. Very large isolate-level or variant-level sheets are represented by compact embedded summaries here, with the complete editable datasets supplied as the matching Excel supplementary tables.",
    )

    add_sheet(doc, "Supplementary Table S1A.", "Locked model policy and endpoint definitions.", "Supplementary_Table_S01_", "model_policy", font_size=5.5)
    add_sheet(doc, "Supplementary Table S1B.", "Public-training apparent performance smoke check.", "Supplementary_Table_S01_", "training_apparent_summary", font_size=5.5)

    add_sheet(doc, "Supplementary Table S2A.", "Locked evaluation summary by cohort and drug.", "Supplementary_Table_S02_", "summary_by_drug", font_size=5.4)
    add_sheet(doc, "Supplementary Table S2B.", "Locked evaluation summary by cohort, drug, and analysis group.", "Supplementary_Table_S02_", "summary_overall", font_size=5.4)

    add_sheet(doc, "Supplementary Table S3A.", "Local paired IPM/MEM phenotype-group metrics.", "Supplementary_Table_S03_", "pair_group_metrics", font_size=5.4)
    add_sheet(doc, "Supplementary Table S3B.", "Local paired IPM/MEM phenotype-group counts.", "Supplementary_Table_S03_", "pair_group_counts", font_size=6.2)

    add_sheet(doc, "Supplementary Table S4.", "First-pass strict mechanism predictability metrics.", "Supplementary_Table_S04_", "strict_predictability", font_size=5.3)

    add_large_table_compact_views(doc)

    add_sheet(doc, "Supplementary Table S7A.", "High-confidence paired IPM/MEM statistical tests.", "Supplementary_Table_S07_", "paired_tests", font_size=5.8)
    add_sheet(doc, "Supplementary Table S7B.", "Global subtype association tests.", "Supplementary_Table_S07_", "global_tests", font_size=6.3)
    add_sheet(doc, "Supplementary Table S7C.", "Predictability summary by high-confidence subtype.", "Supplementary_Table_S07_", "predictability_summary", font_size=5.8)

    add_sheet(doc, "Supplementary Table S8A.", "Local ST counts.", "Supplementary_Table_S08_", "st_counts", font_size=6.2)
    add_sheet(doc, "Supplementary Table S8B.", "Leave-one-ST-out sensitivity analysis.", "Supplementary_Table_S08_", "leave_one_st_out", font_size=6.2)
    add_sheet(doc, "Supplementary Table S8C.", "ST-adjusted model comparison.", "Supplementary_Table_S08_", "st_adjusted_models", font_size=5.7)
    add_sheet(doc, "Supplementary Table S8D.", "Paired error by ST.", "Supplementary_Table_S08_", "paired_error_by_st", font_size=5.8)

    add_sheet(doc, "Supplementary Table S9A.", "Clinical warning-rule table.", "Supplementary_Table_S09_", "clinical_warning_rules", font_size=5.2)
    add_sheet(doc, "Supplementary Table S9B.", "Rule-level performance.", "Supplementary_Table_S09_", "rule_level_performance", font_size=5.4)
    add_sheet(doc, "Supplementary Table S9C.", "MEM warning enrichment tests.", "Supplementary_Table_S09_", "mem_warning_enrichment", font_size=5.5)

    false_cases = read_table("Supplementary_Table_S10_", "false_susceptible_cases")
    false_cols = [
        "drug",
        "genome_id",
        "ipm_mem_pair_group",
        "mechanism_subtype_high_confidence",
        "clinical_rule",
        "actual_mic_log2",
        "pred_mic_log2",
        "signed_log2_error",
        "abs_log2_error",
        "prob_ns",
        "oprd_deep_severity",
        "dacB_severity",
    ]
    add_table_caption(doc, "Supplementary Table S10A.", "False-susceptible cases with selected mechanism context.")
    add_dataframe(doc, selected(false_cases, false_cols), font_size=5.0)

    top_errors = read_table("Supplementary_Table_S10_", "top_error_cases")
    top_cols = [
        "drug",
        "genome_id",
        "mechanism_subtype_high_confidence",
        "mechanism_subtype_strict",
        "actual_mic_log2",
        "pred_mic_log2",
        "signed_log2_error",
        "abs_log2_error",
        "mic_ea_pm1",
        "sns_correct",
        "oprd_deep_severity",
        "dacB_severity",
    ]
    add_table_caption(doc, "Supplementary Table S10B.", "Top absolute MIC error cases with selected mechanism context.")
    add_dataframe(doc, selected(top_errors, top_cols), font_size=5.0)

    for label, text, sheet in [
        ("Supplementary Table S11A.", "Calibration summary by high-confidence subtype.", "calibration_summary"),
        ("Supplementary Table S11B.", "Calibration by probability bin.", "calibration_bins"),
        ("Supplementary Table S11C.", "Susceptible-call risk by confidence.", "susceptible_call_risk"),
        ("Supplementary Table S11D.", "Breakpoint-zone error summary.", "breakpoint_zone_errors"),
        ("Supplementary Table S11E.", "Phenotype-combination paired error summary.", "phenotype_combo_errors"),
    ]:
        add_sheet(doc, label, text, "Supplementary_Table_S11_", sheet, font_size=5.8)

    add_sheet(doc, "Supplementary Table S12A.", "Error taxonomy summary.", "Supplementary_Table_S12_", "error_taxonomy_summary", font_size=5.8)
    add_sheet(doc, "Supplementary Table S12B.", "Safety-gate evaluation summary.", "Supplementary_Table_S12_", "safety_gate_summary", font_size=5.7)


def main() -> None:
    writer = load_writer()
    writer.DATE = DATE
    doc = writer.supplementary_doc()
    make_landscape(doc)
    add_supplementary_tables(doc)
    doc.save(OUT_PATH)
    print(OUT_PATH)


if __name__ == "__main__":
    main()
