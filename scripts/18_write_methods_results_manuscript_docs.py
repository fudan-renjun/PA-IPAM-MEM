#!/usr/bin/env python3
"""Create manuscript Methods/Results and supplementary legend Word drafts."""

from __future__ import annotations

from pathlib import Path

import pandas as pd
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


PROJECT = Path(__file__).resolve().parents[1]
OUT = PROJECT / "投稿" / "manuscript_drafts"
DATE = "2026-06-09"


TITLE = "Mechanism-dependent predictability of imipenem and meropenem MICs in Pseudomonas aeruginosa"

ABSTRACT = (
    "Whole-genome sequencing-based antimicrobial susceptibility prediction is often evaluated as a "
    "single pooled task, although resistance mechanisms can determine whether an MIC estimate is "
    "biologically interpretable or clinically unsafe. We investigated why carbapenem MIC prediction "
    "was poor for Pseudomonas aeruginosa by re-analysing imipenem (IPM) and meropenem (MEM) with a "
    "unified public-only locked model and a mechanism-resolved local validation design. Public data "
    "were used for model development, whereas local clinical genomes and phenotypes were reserved "
    "for locked validation. Local isolates were assigned to strict and high-confidence mechanism "
    "strata using OprD-loss, acquired carbapenemase, efflux-regulator, and AmpC-axis evidence. "
    "Prediction performance was evaluated by essential agreement, absolute and signed log2 MIC "
    "error, categorical agreement, very major error with true non-susceptible isolates as "
    "denominator, paired IPM/MEM error, sequence-type sensitivity, and mechanism-aware warning "
    "rules. The locked local cohort showed marked endpoint divergence: IPM achieved 69.6% "
    "within-one-dilution agreement, MAE 1.22 log2 dilutions, categorical agreement 74.8%, and VME "
    "19.0%, whereas MEM achieved 13.6% within-one-dilution agreement, MAE 2.57, categorical "
    "agreement 44.7%, and VME 83.3%. High-confidence OprD-loss and composite backgrounds were "
    "interpretable for IPM, with high agreement and no observed VME, but generated systematic MEM "
    "underprediction. MEM OprD-loss isolates had 18.2% within-one-dilution agreement, MAE 3.07, "
    "bias -3.07, VME 94.7%, and large underprediction in 77.3%. Paired isolate and sequence-type "
    "analyses indicated that this MEM-worse signal was drug-specific rather than a single-clone "
    "artefact. Mechanism-aware safety gates eliminated released MEM very major errors at the cost "
    "of lower reporting coverage. These findings support drug-specific, mechanism-aware reliability "
    "labels for carbapenem WGS-AST in P. aeruginosa and provide a research-use FASTA prototype for "
    "predictability-stratum assignment."
)

KEYWORDS = [
    "Pseudomonas aeruginosa",
    "carbapenem",
    "imipenem",
    "meropenem",
    "minimum inhibitory concentration",
    "whole-genome sequencing",
    "OprD",
    "antimicrobial susceptibility prediction",
]


def setup_doc() -> Document:
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(0.85)
    section.right_margin = Inches(0.85)
    styles = doc.styles
    styles["Normal"].font.name = "Arial"
    styles["Normal"].font.size = Pt(10.5)
    styles["Title"].font.name = "Arial"
    styles["Title"].font.size = Pt(16)
    styles["Heading 1"].font.name = "Arial"
    styles["Heading 1"].font.size = Pt(14)
    styles["Heading 2"].font.name = "Arial"
    styles["Heading 2"].font.size = Pt(12)
    return doc


def add_title(doc: Document, title: str, subtitle: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(title)
    run.bold = True
    run.font.size = Pt(16)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(subtitle)
    run.italic = True
    run.font.size = Pt(10.5)
    doc.add_paragraph()


def add_p(doc: Document, text: str) -> None:
    p = doc.add_paragraph(text)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.08


def add_bullet(doc: Document, text: str) -> None:
    p = doc.add_paragraph(text, style="List Bullet")
    p.paragraph_format.space_after = Pt(2)


def add_caption(doc: Document, label: str, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(8)
    r = p.add_run(label)
    r.bold = True
    p.add_run(" " + text)


def add_abstract_and_keywords(doc: Document) -> None:
    doc.add_heading("Abstract", level=1)
    add_p(doc, ABSTRACT)
    doc.add_heading("Keywords", level=1)
    add_p(doc, "; ".join(KEYWORDS))


def find_table_file(pattern: str) -> Path:
    matches = sorted(PROJECT.glob(f"*/manuscript_tables/{pattern}"))
    if not matches:
        raise FileNotFoundError(f"Could not find manuscript table file matching {pattern}")
    return matches[0]


def fmt_cell(value: object) -> str:
    if pd.isna(value):
        return ""
    if isinstance(value, float):
        if value.is_integer():
            return str(int(value))
        return f"{value:.3g}"
    return str(value)


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_margins(cell, top: int = 60, start: int = 60, bottom: int = 60, end: int = 60) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in [("top", top), ("start", start), ("bottom", bottom), ("end", end)]:
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def style_cell(cell, *, font_size: float, bold: bool = False, align: WD_ALIGN_PARAGRAPH = WD_ALIGN_PARAGRAPH.CENTER) -> None:
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    set_cell_margins(cell)
    for paragraph in cell.paragraphs:
        paragraph.alignment = align
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.paragraph_format.line_spacing = 1.0
        for run in paragraph.runs:
            run.bold = bold
            run.font.name = "Arial"
            run.font.size = Pt(font_size)


def is_text_heavy_column(col: str) -> bool:
    text_tokens = [
        "cohort",
        "subtype",
        "rule",
        "rationale",
        "drug",
        "mechanism",
        "interpretation",
    ]
    lowered = str(col).lower()
    return any(token in lowered for token in text_tokens)


def add_dataframe_table(doc: Document, df: pd.DataFrame, max_font_size: float = 7.0) -> None:
    table = doc.add_table(rows=1, cols=len(df.columns))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    header = table.rows[0].cells
    for j, col in enumerate(df.columns):
        header[j].text = str(col)
        set_cell_shading(header[j], "D9EAF7")
        style_cell(header[j], font_size=max_font_size, bold=True)
    for _, row in df.iterrows():
        table_row = table.add_row()
        cells = table_row.cells
        row_idx = len(table.rows) - 1
        for j, col in enumerate(df.columns):
            cells[j].text = fmt_cell(row[col])
            if row_idx % 2 == 0:
                set_cell_shading(cells[j], "F7FBFD")
            align = WD_ALIGN_PARAGRAPH.LEFT if is_text_heavy_column(str(col)) else WD_ALIGN_PARAGRAPH.CENTER
            style_cell(cells[j], font_size=max_font_size, align=align)
    doc.add_paragraph()


def add_main_tables(doc: Document) -> None:
    table1 = find_table_file("main_tables/Table_1_cohort_composition_and_locked_model_evaluation_2026-06-07.xlsx")
    table2 = find_table_file("main_tables/Table_2_high_confidence_clinical_interpretation_rules_2026-06-07.xlsx")

    doc.add_heading("Main Tables", level=1)
    add_caption(
        doc,
        "Table 1. Cohort composition and locked model evaluation.",
        "A, Cohort composition and phenotype availability. B, Locked evaluation performance under "
        "the unified public-only IPM/MEM model. VME and ME use category-specific denominators."
    )
    add_p(doc, "Table 1A. Cohort composition and phenotype availability.")
    add_dataframe_table(doc, pd.read_excel(table1, sheet_name="cohort_composition_display"))
    add_p(doc, "Table 1B. Locked model evaluation.")
    add_dataframe_table(doc, pd.read_excel(table1, sheet_name="locked_evaluation_display"), max_font_size=6.5)

    add_caption(
        doc,
        "Table 2. High-confidence mechanism strata and clinical interpretation rules.",
        "IPM/MEM performance by high-confidence mechanism subtype and study-defined interpretation "
        "rule. VME is false-susceptible calls divided by true non-susceptible isolates."
    )
    add_dataframe_table(doc, pd.read_excel(table2, sheet_name="clinical_rules_display"), max_font_size=5.8)


def add_supplementary_table_file_index(doc: Document) -> None:
    files = sorted(PROJECT.glob("*/manuscript_tables/supplementary_tables/Supplementary_Table_S*.xlsx"))
    if not files:
        return
    rows = []
    for file in files:
        name = file.name
        label = name.split("_")[2] if len(name.split("_")) > 2 else ""
        try:
            xl = pd.ExcelFile(file)
            sheet_names = xl.sheet_names
            sheet_count = len(sheet_names)
            sheet_dims = []
            for sheet in sheet_names:
                frame = pd.read_excel(file, sheet_name=sheet)
                sheet_dims.append(f"{sheet}: {frame.shape[0]} rows x {frame.shape[1]} columns")
            sheet_summary = "; ".join(sheet_dims)
        except Exception:
            sheet_count = ""
            sheet_summary = ""
        rows.append(
            {
                "Supplementary table": label.replace("S", "S"),
                "Excel file": name,
                "Sheets": sheet_count,
                "Sheet contents": sheet_summary,
            }
        )
    add_p(
        doc,
        "The full supplementary tables are provided as editable Excel files because several contain "
        "isolate-level or variant-level records that are too large for readable Word layout."
    )
    add_dataframe_table(doc, pd.DataFrame(rows), max_font_size=7.0)


def main_manuscript() -> Document:
    doc = setup_doc()
    add_title(doc, TITLE, f"Full manuscript draft with Introduction, Methods, Results, Discussion, and legends; generated {DATE}")
    add_abstract_and_keywords(doc)

    doc.add_heading("Introduction", level=1)
    add_p(
        doc,
        "Whole-genome sequencing (WGS) is increasingly used to support antimicrobial resistance "
        "surveillance and to develop genotype-based antimicrobial susceptibility testing (AST) "
        "models. For organisms with relatively discrete resistance determinants, WGS-based models "
        "can provide clinically useful categorical predictions. Pseudomonas aeruginosa, however, "
        "presents a more difficult problem. Carbapenem susceptibility in this species is shaped by "
        "multiple interacting mechanisms, including OprD porin loss, acquired carbapenemases, "
        "AmpC-associated regulation, multidrug efflux systems, and background-dependent or "
        "expression-dependent effects. As a result, a single WGS-AST model can appear to perform "
        "well in some isolates and poorly in others, even when the same drug and the same modelling "
        "policy are used."
    )
    add_p(
        doc,
        "This issue is especially important for imipenem (IPM) and meropenem (MEM). Both are "
        "carbapenems and are often discussed together, yet their genotype-phenotype relationships "
        "are not necessarily interchangeable. Severe oprD disruption is biologically central to IPM "
        "resistance and may create a more discrete genomic signal for IPM prediction. In contrast, "
        "MEM MICs can be influenced by a broader combination of OprD, efflux, AmpC-axis, and other "
        "mechanisms, making MEM prediction more vulnerable to missing expression, regulatory, or "
        "context-dependent effects. A model that is accurate for IPM in OprD-dominant backgrounds "
        "may therefore still underpredict MEM in the same isolates."
    )
    add_p(
        doc,
        "Most previous WGS or machine-learning studies of P. aeruginosa carbapenem resistance have "
        "asked whether resistance or susceptibility can be predicted from genomic data. That framing "
        "is necessary but incomplete. Clinical use of WGS-AST requires knowing when a prediction is "
        "reliable, when it is directionally informative but not MIC-accurate, and when it should not "
        "be used as a stand-alone call. This is particularly relevant for MIC-level prediction, where "
        "errors of several two-fold dilutions may lead to false-susceptible reports far above a "
        "breakpoint. Thus, prediction failure should not be treated only as a modelling limitation; "
        "it can also be analysed as a biological and translational signal."
    )
    add_p(
        doc,
        "We hypothesised that carbapenem MIC predictability in P. aeruginosa is mechanism-dependent. "
        "Under this hypothesis, isolates with discrete, high-effect genomic mechanisms, such as "
        "severe OprD loss, should show more predictable MIC behaviour for the drug most tightly "
        "linked to that mechanism. By contrast, isolates whose resistance depends on regulatory, "
        "expression-dependent, or unresolved mechanisms should show larger MIC error and higher "
        "categorical safety risk. We further hypothesised that IPM and MEM would differ in their "
        "mechanism-specific predictability despite belonging to the same antibiotic class."
    )
    add_p(
        doc,
        "To test these hypotheses, we rebuilt a study-specific unified public-training-only IPM/MEM "
        "prediction layer using the same model recipe for both drugs, joined locked prediction "
        "errors to genome-defined mechanism annotations, and quantified subtype-specific MIC and "
        "categorical errors in locked public external and local clinical validation cohorts. We then "
        "performed paired within-isolate IPM/MEM comparisons, sequence-type sensitivity analyses, "
        "and clinical warning-rule analyses. Finally, we implemented a research-use FASTA upload "
        "prototype to demonstrate how mechanism-aware predictability flags could accompany numerical "
        "IPM/MEM MIC estimates."
    )

    doc.add_heading("Methods", level=1)
    doc.add_heading("Study design and analysis overview", level=2)
    add_p(
        doc,
        "We designed this study as a mechanism-resolved follow-up to a whole-genome sequence "
        "(WGS)-based minimum inhibitory concentration (MIC) prediction framework in Pseudomonas "
        "aeruginosa. The motivating observation was that carbapenem MIC prediction, particularly "
        "for imipenem (IPM) and meropenem (MEM), was heterogeneous across isolates and appeared to "
        "fail in a drug-specific manner. We therefore asked whether genome-defined resistance "
        "mechanism backgrounds could distinguish isolates for which carbapenem MICs were predictable "
        "from those for which numerical MIC or susceptible/non-susceptible calls were unreliable."
    )
    add_p(
        doc,
        "The analysis used a pre-specified public-training-only modelling layer and two locked "
        "evaluation cohorts. Public training data were used only for model fitting. Locked public "
        "external genomes and a locked local clinical validation cohort were not used for model "
        "training, model selection, threshold optimisation, or mechanism-rule selection. The local "
        "validation cohort provided the primary paired IPM/MEM analysis because both drug phenotypes, "
        "assembly data, mechanism annotations, and local clinical metadata were available."
    )

    doc.add_heading("Cohorts and phenotypic endpoints", level=2)
    add_p(
        doc,
        "The public training cohort contained 3,138 P. aeruginosa genomes/features, including 312 "
        "IPM and 3,065 MEM MIC/SIR records. Locked public external validation contained 637 genomes, "
        "with 5 IPM and 637 MEM phenotypes. The local clinical validation cohort contained 147 "
        "isolates/assemblies, with 115 IPM and 103 MEM phenotypes. IPM and MEM MICs were analysed "
        "on a log2 dilution scale. Categorical interpretation was expressed as susceptible (S) or "
        "non-susceptible (NS), using the locked endpoint-specific breakpoints embedded in the model "
        "policy: IPM log2 breakpoint 2.0 and MEM log2 breakpoint 1.0."
    )
    add_p(
        doc,
        "For local isolates, paired phenotype groups were defined according to available IPM and MEM "
        "phenotypes: IPM-R/MEM-R, IPM-R/MEM-S or I, IPM-S/MEM-R, IPM-S/MEM-S, and IPM/MEM incomplete. "
        "The paired IPM/MEM analysis was restricted to isolates with prediction-error rows for both "
        "drugs."
    )

    doc.add_heading("Unified public-only IPM/MEM prediction model", level=2)
    add_p(
        doc,
        "To avoid confounding downstream mechanism comparisons by endpoint-specific model choices, "
        "IPM and MEM were re-modelled using the same locked recipe. For each endpoint, a public-only "
        "histogram gradient boosting gate estimated the probability of non-susceptibility, followed "
        "by histogram gradient boosting MIC regression. Both endpoints used all numeric features, a "
        "hard-gate threshold of 0.45, snapping to training MIC levels, and no post-hoc cap policy. "
        "This shared policy generated a study-specific unified IPM/MEM prediction-error table."
    )
    add_p(
        doc,
        "Predicted MICs were compared against observed MICs on the log2 scale. Primary continuous "
        "metrics were essential agreement within +/-1 two-fold dilution, mean absolute error (MAE), "
        "signed bias, exact agreement within +/-0.5 log2 dilution, and large error >2 log2 dilutions. "
        "Categorical metrics included categorical agreement, major error (ME; predicted NS among "
        "true S isolates), very major error (VME; predicted S among true NS isolates), and gate AUC "
        "where both S and NS isolates were present. ME and VME were calculated using standard "
        "category-specific denominators throughout the manuscript: ME = false resistant calls / true "
        "S isolates, and VME = false susceptible calls / true NS isolates. When an analysis used the "
        "proportion of false-susceptible calls among all evaluated rows, this was reported separately "
        "as a false-susceptible all-row proportion and was not labelled as VME."
    )
    add_p(
        doc,
        "Because IPM and MEM have different locked breakpoints in the model policy, categorical "
        "safety metrics were interpreted as breakpoint-dependent clinical summaries. The primary "
        "drug-comparison evidence therefore emphasised breakpoint-independent MIC-scale metrics "
        "including within-one-dilution agreement, MAE, signed bias, and paired absolute error. "
        "VME and ME were used to quantify clinical reporting risk under the specified breakpoint "
        "policy rather than to define the biological mechanism effect."
    )

    doc.add_heading("Genome-defined mechanism annotation", level=2)
    add_p(
        doc,
        "Mechanism annotation was constructed in two layers. A first-pass strict layer assigned "
        "isolates to OprD-loss, acquired carbapenemase-mediated, AmpC-associated genotype, "
        "efflux-associated genotype, composite, or no strict mechanism classes. OprD-loss required "
        "severe oprD evidence, including oprD disruption, truncation, high-confidence disruptive "
        "effect, >=10 bp indel, or OprD length <430 amino acids. Strict acquired carbapenemase "
        "evidence included VIM, IMP, NDM, KPC, SPM, GIM, or carbapenemase-associated GES variants. "
        "AmpC-associated genotype required strict ampC/ampR/dacB/ampD mutation evidence. "
        "Efflux-associated genotype required strict regulator mutations in mexR, nalD, nfxB, mexS, "
        "or mexZ. Composite subtypes contained at least two strict mechanism classes."
    )
    add_p(
        doc,
        "Because PAO1-relative missense variation was widespread in several resistance-associated "
        "genes, broad missense calls were not treated as manuscript-grade mechanism drivers. We "
        "therefore derived a high-confidence evidence layer that prioritised disruptive or severe "
        "events and separated strict driver genes from broad context variation. High-confidence "
        "subtypes included OprD-loss, high-confidence composite, AmpC-axis disruptive, and no "
        "high-confidence driver. Efflux and AmpC calls were interpreted as genotype-associated "
        "signals; no claim of overexpression or functional derepression was made without expression "
        "or functional data."
    )

    doc.add_heading("Paired IPM/MEM analysis and ST sensitivity", level=2)
    add_p(
        doc,
        "To determine whether MEM failure reflected a drug-specific pattern within the same genomes, "
        "we compared IPM and MEM absolute MIC errors in paired local isolates. MEM-minus-IPM absolute "
        "error was summarised overall, within strict mechanism subtypes, within high-confidence "
        "subtypes, and within IPM/MEM phenotype groups. Paired Wilcoxon tests compared IPM and MEM "
        "absolute errors; exact binomial tests evaluated whether the fraction of isolates with "
        "larger MEM than IPM error exceeded 50%."
    )
    add_p(
        doc,
        "MLST sequence types (STs) were called using exact allele-profile matching. ST sensitivity "
        "analyses included ST count summaries, paired MEM-minus-IPM error summaries by ST, "
        "leave-one-major-ST-out analyses, and exploratory linear models comparing mechanism subtype, "
        "ST adjustment group, and combined mechanism-plus-ST models. These analyses were used as "
        "sensitivity evidence rather than causal lineage adjustment because most STs were sparse."
    )

    doc.add_heading("Clinical warning rules and prototype web implementation", level=2)
    add_p(
        doc,
        "Mechanism-specific performance summaries were translated into clinical interpretation "
        "rules. Strata with high essential agreement and no categorical safety errors were labelled "
        "trustworthy or exploratory trustworthy depending on sample size. Strata with high VME, "
        "large negative bias, high MAE, or poor essential agreement were labelled warning or "
        "exploratory warning. Strata without a high-confidence driver were labelled "
        "mechanism-unresolved warning, reflecting the need for phenotype confirmation rather than "
        "absence of mechanism."
    )
    add_p(
        doc,
        "As a companion resource, we implemented a research-use FASTA upload prototype. The tool "
        "computes basic assembly QC, screens selected PAO1 target loci using k-mer coverage, assigns "
        "a mechanism-informed predictability stratum, calls the bundled locked public-only IPM/MEM "
        "model to generate FASTA-derived prototype MIC estimates, and reports stratum-specific "
        "reliability flags. The web prototype is not a validated clinical antimicrobial "
        "susceptibility testing system and is distinct from the full feature-generation pipeline "
        "used for manuscript evaluation."
    )

    doc.add_heading("Statistical analysis", level=2)
    add_p(
        doc,
        "Subtype differences in absolute log2 error were assessed using Kruskal-Wallis tests. "
        "Subtype differences in within-one-dilution agreement were assessed using chi-square tests. "
        "Paired IPM/MEM error differences were assessed using paired Wilcoxon tests and exact "
        "binomial tests. Fisher exact tests were used for enrichment analyses of large MEM "
        "underprediction among warning strata. Standard VME was retained as a clinical safety "
        "outcome, but was not interpreted as an OprD-enrichment signal when rates were broadly "
        "high across true non-susceptible MEM isolates. Calibration summaries included mean predicted "
        "probability of non-susceptibility, observed NS rate, Brier score, and bin-based expected "
        "calibration error. All analyses were conducted using locked model predictions and "
        "pre-defined validation partitions."
    )

    doc.add_heading("Results", level=1)
    doc.add_heading("A unified public-only model exposed a drug-specific carbapenem predictability problem", level=2)
    add_p(
        doc,
        "We first re-established a common IPM/MEM prediction-error layer using the same public-only "
        "model recipe for both drugs (Figure 1; Table 1). In the locked local clinical cohort, IPM "
        "prediction was substantially more accurate than MEM prediction. IPM achieved 69.6% "
        "essential agreement within +/-1 dilution, MAE 1.22 log2 dilutions, bias +0.49 log2 "
        "dilutions, categorical agreement 74.8%, VME 19.0%, and gate AUC 0.80 across 115 local IPM "
        "rows. In contrast, MEM achieved only 13.6% essential agreement, MAE 2.57 log2 dilutions, "
        "bias -1.44 log2 dilutions, categorical agreement 44.7%, VME 83.3%, and gate AUC 0.55 "
        "across 103 local MEM rows (Figure 2B; Table 1)."
    )
    add_p(
        doc,
        "This divergence was not a trivial consequence of modelling different endpoints with "
        "different policies, because both endpoints used the same feature policy, gate threshold, "
        "regression architecture, snapping rule, and cap policy. The locked public external cohort "
        "provided broader MEM context, where MEM performance was higher than in the local cohort "
        "(637 rows; 56.4% within +/-1 dilution; MAE 1.37; categorical agreement 80.1%; VME 15.6%), "
        "but the local cohort revealed a clinically important pattern of MEM underprediction that "
        "motivated mechanism-resolved analysis."
    )
    add_p(
        doc,
        "The marked difference between local MEM and locked public external MEM performance was "
        "therefore treated as a feature-shift signal rather than ignored. Several non-exclusive "
        "explanations are plausible: the local cohort was enriched for difficult clinical "
        "carbapenem-resistant backgrounds, contained paired IPM/MEM phenotypes that exposed "
        "discordant drug behaviour, and included many high-MIC or censored non-susceptible cases in "
        "which MEM was systematically underpredicted. By contrast, the locked public external MEM "
        "set was larger and more heterogeneous, and its aggregate performance averaged across "
        "mechanism backgrounds with lower local-like warning enrichment. Because the local MEM "
        "failure was also observed in breakpoint-independent metrics, including within-one-dilution "
        "agreement, MAE, bias, and paired MEM-minus-IPM absolute error, the central conclusion does "
        "not rely solely on categorical VME."
    )
    add_p(
        doc,
        "Among local isolates, paired IPM/MEM phenotype groups included IPM-R/MEM-R, IPM-R/MEM-S "
        "or I, IPM-S/MEM-R, IPM-S/MEM-S, and incomplete groups (Figure 2C). This paired phenotype "
        "structure enabled direct within-isolate comparisons of drug-specific prediction error."
    )

    doc.add_heading("Genome-defined mechanism subtypes revealed heterogeneous carbapenem backgrounds", level=2)
    add_p(
        doc,
        "The local cohort contained substantial mechanism heterogeneity (Figure 2D-G). In the first-pass "
        "strict mechanism layer, 31 isolates were classified as composite, 47 as efflux-associated "
        "genotype, 49 as no strict mechanism, and 20 as OprD-loss. Composite local isolates all "
        "carried severe OprD-loss evidence and almost all carried strict efflux-regulator evidence, "
        "whereas OprD-loss isolates were defined by severe OprD loss without another strict driver. "
        "Locked public external MEM isolates were enriched for OprD-loss and composite backgrounds, "
        "including 422 composite and 190 OprD-loss isolates."
    )
    add_p(
        doc,
        "Broad PAO1-relative variation was common across resistance-associated target genes. For this "
        "reason, broad missense calls were retained as context but not used as sole manuscript-grade "
        "mechanism drivers. The high-confidence layer therefore prioritised disruptive OprD events, "
        "high-confidence composite backgrounds, and disruptive AmpC-axis signals. This two-layer "
        "framework allowed the analysis to separate biological mechanism evidence from phylogenetic "
        "or background polymorphism (Figure 2F,G)."
    )

    doc.add_heading("High-confidence mechanism strata separated IPM-predictable from MEM-warning backgrounds", level=2)
    add_p(
        doc,
        "Mechanism-resolved evaluation showed that IPM and MEM were not interchangeable as "
        "carbapenem endpoints (Figure 3; Table 2). In the strict local mechanism analysis, IPM was "
        "highly predictable in OprD-loss and composite backgrounds. IPM OprD-loss isolates achieved "
        "93.3% within-one-dilution agreement, MAE 0.34, bias +0.34, categorical agreement 100%, and "
        "VME 0%. IPM composite isolates achieved 90.5% within-one-dilution agreement, MAE 0.39, "
        "bias +0.39, categorical agreement 100%, and VME 0%."
    )
    add_p(
        doc,
        "MEM behaved in the opposite direction in the same broad mechanism space. MEM OprD-loss "
        "isolates achieved only 25.0% within-one-dilution agreement, MAE 2.75, bias -2.75, "
        "categorical agreement 25.0%, and VME 100.0% in the strict local analysis. MEM composite "
        "isolates achieved 5.3% within-one-dilution agreement, MAE 3.05, bias -3.05, categorical "
        "agreement 31.6%, and VME 68.4%. Thus, the same OprD/composite backgrounds that were "
        "interpretable for IPM were warning backgrounds for MEM."
    )
    add_p(
        doc,
        "The high-confidence layer strengthened this conclusion. IPM OprD-loss isolates (n=28) "
        "showed 89.3% within-one-dilution agreement, MAE 0.38, categorical agreement 100%, VME 0%, "
        "and ME 0%. IPM high-confidence composite isolates (n=8) and AmpC-axis disruptive isolates "
        "(n=3) each showed 100% within-one-dilution agreement and no categorical safety errors, but "
        "were treated as exploratory because of sample size. By contrast, MEM OprD-loss isolates "
        "(n=22) showed 18.2% within-one-dilution agreement, MAE 3.07, bias -3.07, categorical "
        "agreement 18.2%, VME 94.7%, and large error >2 log2 dilutions in 77.3%. MEM high-confidence "
        "composite and AmpC-axis disruptive groups showed the same direction of failure, with poor "
        "essential agreement and negative bias."
    )
    add_p(
        doc,
        "Observed-versus-predicted MIC summaries showed that this was a directional prediction "
        "failure rather than random scatter. In MEM OprD-loss isolates, the median observed MIC was "
        "3.5 log2 dilutions whereas the median predicted MIC was 0.0 log2 dilutions, and 77.3% of "
        "isolates had large underprediction >2 log2 dilutions. MEM high-confidence composite "
        "isolates also showed a downward predicted-versus-observed MIC shift, whereas high-confidence "
        "IPM OprD-loss and composite backgrounds remained tightly aligned with observed MICs "
        "(Figure 3E,F)."
    )
    add_p(
        doc,
        "Subtype association tests supported mechanism-dependent IPM predictability. For IPM, "
        "absolute error differed across strict subtypes (Kruskal-Wallis p=0.0023), and "
        "within-one-dilution agreement differed across subtypes (chi-square p=0.00443). MEM subtype "
        "differences were borderline because nearly all local MEM subtype groups performed poorly "
        "(Kruskal-Wallis p=0.0557; chi-square p=0.0562)."
    )

    doc.add_heading("Paired isolate analysis confirmed that MEM underprediction was drug-specific and not driven by a single ST", level=2)
    add_p(
        doc,
        "Within the 101 local paired IPM/MEM prediction rows, MEM absolute error exceeded IPM "
        "absolute error in 72 isolates (71.3%; binomial p=2.24e-05), and paired absolute errors were "
        "significantly larger for MEM than IPM (paired Wilcoxon p=1.13e-07; Figure 4A,B). The paired "
        "difference was most pronounced in composite and OprD-loss backgrounds. In strict composite "
        "isolates, MEM was worse than IPM in 18/18 paired isolates, with mean MEM-minus-IPM absolute "
        "error +2.64 log2 dilutions and paired Wilcoxon p=0.000167. In strict OprD-loss isolates, "
        "MEM was worse in 12/12 paired isolates, with mean MEM-minus-IPM absolute error +2.48 log2 "
        "dilutions and paired Wilcoxon p=0.000488."
    )
    add_p(
        doc,
        "The MEM-worse-than-IPM signal persisted after lineage sensitivity analyses. ST244 was the "
        "largest local ST in the prediction-analysis subset, but contained only 12 paired isolates. "
        "After excluding ST244, 89 paired isolates remained and the MEM-worse fraction was 73.0%, "
        "with paired Wilcoxon p=4.49e-07. Excluding ST-unassigned isolates or ST463 similarly "
        "preserved the signal. In exploratory ST-adjusted models, mechanism subtype explained more "
        "variation in MEM-minus-IPM error than ST group alone (R2 0.140 for mechanism subtype, 0.011 "
        "for ST group, and 0.193 for the combined mechanism-plus-ST model; Figure 4C-E). ST-level "
        "paired summaries further showed positive MEM-minus-IPM absolute error across multiple STs, "
        "supporting a non-clonal MEM-worse signal rather than a single-lineage artefact (Figure 4F)."
    )

    doc.add_heading("Mechanism-aware warning rules identified MEM false-susceptible risk far above the breakpoint", level=2)
    add_p(
        doc,
        "We next translated mechanism-specific performance into clinical interpretation rules "
        "(Figure 5; Table 2). IPM OprD-loss was classified as trustworthy because it combined high "
        "essential agreement with no categorical safety errors. IPM high-confidence composite and "
        "AmpC-axis disruptive groups were classified as exploratory trustworthy because performance "
        "was strong but sample sizes were small. In contrast, MEM OprD-loss was classified as a "
        "warning stratum because of high VME, high MAE, and large negative bias. MEM high-confidence "
        "composite and AmpC-axis disruptive groups were classified as exploratory warning strata."
    )
    add_p(
        doc,
        "MEM susceptible calls were particularly unsafe in OprD-loss and high-confidence composite "
        "backgrounds. In the MEM OprD-loss group, 19 confident susceptible predictions contained "
        "16 false-susceptible calls (84.2%), with 15 far-from-breakpoint VMEs. High-confidence "
        "composite confident or low-margin susceptible MEM calls were all false susceptible in this "
        "local validation set. Far-from-breakpoint analyses showed that MEM OprD-loss isolates above "
        "the breakpoint had only 5.6% within-one-dilution agreement, MAE 3.47, bias -3.47, VME "
        "94.4%, and large underprediction >2 log2 dilutions in 94.4%."
    )
    add_p(
        doc,
        "In contrast to the all-row false-susceptible proportion, standard MEM VME did not enrich "
        "specifically within OprD-flagged isolates once the denominator was restricted to true "
        "non-susceptible isolates; OprD-flagged and unflagged MEM backgrounds both showed high "
        "VME. The discriminating signal in Figure 5C was therefore large MIC underprediction, "
        "which was enriched in OprD-disruptive backgrounds and provides the quantitative basis "
        "for the warning interpretation."
    )
    add_p(
        doc,
        "A proof-of-concept safety gate that withheld MEM susceptible calls in mechanism-warning or "
        "low-margin contexts released only 13/103 MEM calls, but eliminated all 55 original MEM VMEs "
        "among released predictions. Across all IPM/MEM local predictions, the same gate released "
        "91/218 predictions and avoided all 70 original VMEs among released calls. These results "
        "support reporting mechanism-aware reliability flags rather than unqualified numerical MIC "
        "or susceptible calls in warning strata."
    )

    doc.add_heading("A FASTA-based prototype translated the framework into a reusable predictability report", level=2)
    add_p(
        doc,
        "Finally, we implemented a research-use web prototype to demonstrate how the framework can be "
        "used for new assemblies (Figure 6). Users upload an assembly FASTA file; the application "
        "computes assembly QC, screens selected PAO1 target loci by k-mer coverage, assigns a "
        "mechanism-informed predictability stratum, generates FASTA-derived prototype IPM/MEM MIC "
        "estimates using the bundled locked public-only model artifact, and reports drug-specific "
        "reliability flags. The output explicitly distinguishes interpretable, caution, and "
        "high-risk false-susceptible patterns. The tool is intended as a transparent companion "
        "demonstrator of the study framework, not as a validated clinical AST system."
    )

    doc.add_heading("Discussion", level=1)
    add_p(
        doc,
        "This study shows that carbapenem MIC prediction in P. aeruginosa is not a single endpoint-"
        "level property of a WGS model. Instead, predictability depends on the genomic mechanism "
        "background and on the specific carbapenem being interpreted. Using a unified public-only "
        "model policy for IPM and MEM, we found that high-confidence OprD-loss and composite "
        "backgrounds were highly interpretable for IPM but were warning backgrounds for MEM. The "
        "finding was strongest in paired isolates, where MEM errors exceeded IPM errors within the "
        "same genomes, particularly in OprD-loss and composite strata. This paired design is "
        "important because it separates drug-specific predictability from isolate-level data quality "
        "or cohort composition."
    )
    add_p(
        doc,
        "The most clinically relevant observation was not simply that MEM had poorer average "
        "performance. Rather, MEM errors in warning strata were frequently false-susceptible and "
        "often occurred far above the breakpoint. In the local validation set, MEM OprD-loss "
        "isolates showed large negative bias, poor essential agreement, and a high VME rate, whereas "
        "IPM in the corresponding high-confidence OprD-loss background showed high essential "
        "agreement and no categorical safety errors. These results indicate that WGS-AST outputs for "
        "carbapenems should not be reported as unqualified numerical predictions. They require "
        "drug-specific, mechanism-aware reliability labels."
    )
    add_p(
        doc,
        "The contrast between local and public external MEM performance deserves careful "
        "interpretation. We do not interpret the local MEM result as proof that MEM prediction will "
        "always fail at the same rate in every setting. Instead, it shows that a model with acceptable "
        "aggregate external MEM performance can become unsafe in a clinically enriched local context "
        "containing OprD-loss, composite, and high-MIC non-susceptible backgrounds. This is precisely "
        "the setting in which mechanism-stratified evaluation is needed: aggregate external metrics "
        "can coexist with high local false-susceptible risk in specific mechanism strata."
    )
    add_p(
        doc,
        "Breakpoint asymmetry also matters. The locked MEM breakpoint was lower than the IPM "
        "breakpoint, and this can increase the probability that numerical underprediction crosses "
        "the susceptible/non-susceptible threshold. For that reason, our clinical safety narrative "
        "distinguishes MIC-scale evidence from breakpoint-dependent categorical error. The MIC-scale "
        "evidence was consistent with true MEM difficulty: MEM had much lower within-one-dilution "
        "agreement, larger MAE, negative bias, and larger paired absolute error than IPM in the same "
        "isolates. The VME analysis then quantifies the clinical consequence of this underprediction "
        "under the specified breakpoint policy."
    )
    add_p(
        doc,
        "Biologically, the results are consistent with the idea that severe OprD loss creates a "
        "strong and relatively discrete genomic signal for IPM resistance, while MEM MICs are more "
        "sensitive to additional or unmeasured contributors. Efflux regulation, AmpC-axis variation, "
        "carbapenemase context, and background genetic effects may all modify MEM MICs. Because WGS "
        "alone does not measure expression, porin abundance, enzyme expression, or regulatory state, "
        "genotype-only MEM prediction may fail even when important target genes are annotated. This "
        "does not imply that WGS is uninformative for MEM; rather, it implies that WGS-derived MEM "
        "calls should be interpreted through a mechanism-specific uncertainty framework."
    )
    add_p(
        doc,
        "Our two-layer mechanism framework was essential for this interpretation. Broad PAO1-relative "
        "missense variation was common across several resistance-associated loci and could not be "
        "treated as direct evidence of a causal resistance mechanism. If broad variation were used "
        "alone, many isolates would collapse into composite or regulatory categories, obscuring the "
        "relationship between high-confidence mechanism evidence and prediction error. By separating "
        "broad context from high-confidence disruptive evidence, the analysis identified a clearer "
        "IPM-predictable/OprD-dominant stratum and a MEM-warning counterpart."
    )
    add_p(
        doc,
        "The lineage analyses suggest that the paired IPM/MEM error gap was not explained by a single "
        "dominant ST. ST244 was the largest local ST in the prediction-analysis subset, but removing "
        "it did not attenuate the MEM-worse-than-IPM signal. Exploratory ST-adjusted models also "
        "showed that mechanism subtype explained more MEM-minus-IPM error variation than ST group "
        "alone. These results do not eliminate lineage effects, and the sparse ST distribution limits "
        "formal adjustment, but they support the conclusion that the observed predictability pattern "
        "is not merely a clone-specific artefact."
    )
    add_p(
        doc,
        "The clinical warning-rule analysis illustrates how mechanism-resolved prediction evaluation "
        "can be translated into safer reporting logic. In the local validation set, withholding MEM "
        "susceptible calls in mechanism-warning or low-margin contexts avoided all original MEM VMEs "
        "among released predictions, albeit with low release coverage. This trade-off is expected for "
        "a safety-oriented gate. The goal is not to maximise the number of automated reports, but to "
        "identify contexts in which a model-generated susceptible result should trigger phenotype "
        "confirmation or be reported with explicit caution."
    )
    add_p(
        doc,
        "These findings also have implications for how WGS-AST models should be benchmarked. Aggregate "
        "accuracy can mask opposing behaviours across biological strata. A model can be useful for "
        "one mechanism-drug pairing and unsafe for another, even within the same antibiotic class. "
        "Future WGS-AST evaluations in complex organisms should therefore include mechanism-stratified "
        "MIC error, bias, VME, and breakpoint-distance analyses, rather than relying only on pooled "
        "categorical agreement or global AUC."
    )
    add_p(
        doc,
        "The companion FASTA upload prototype was developed to make this interpretive framework "
        "transparent and reusable. It assigns uploaded assemblies to mechanism-informed "
        "predictability strata, generates FASTA-derived prototype IPM/MEM MIC estimates with the "
        "bundled locked model artifact, and reports reliability flags. This implementation is a "
        "research demonstrator, not a validated diagnostic tool. It does not replace phenotypic AST "
        "or the full feature-generation pipeline, and its selected target-locus k-mer screen cannot "
        "capture all functionally relevant point mutations, expression changes, or acquired "
        "resistance genes. Its purpose is to demonstrate how numerical WGS-AST outputs can be paired "
        "with mechanism-aware caution labels."
    )
    add_p(
        doc,
        "The study has limitations. First, the primary discovery and paired IPM/MEM analyses were "
        "performed in a single local clinical validation cohort, and some high-confidence strata, "
        "particularly composite and AmpC-axis disruptive categories, had small sample sizes. These "
        "groups were therefore labelled exploratory. Second, the locked public external cohort was "
        "informative mainly for MEM because public IPM phenotypes were sparse. Third, AmpC and efflux "
        "mechanisms were inferred from genomic evidence and were not validated by expression or "
        "functional assays. Fourth, MIC testing itself has inherent dilution-level variability, so "
        "small near-breakpoint differences should be interpreted cautiously. Finally, the web "
        "prototype uses a simplified FASTA-derived feature representation and should not be "
        "confused with the full manuscript evaluation pipeline."
    )
    add_p(
        doc,
        "In conclusion, WGS-based carbapenem MIC prediction in P. aeruginosa is mechanism- and "
        "drug-dependent. High-confidence OprD-loss and composite backgrounds can be interpretable "
        "for IPM, but the same or related backgrounds can generate dangerous MEM underprediction. "
        "Mechanism-aware reliability labels and safety gates may therefore be necessary for the "
        "clinical translation of WGS-AST in P. aeruginosa. Rather than asking only whether a model "
        "predicts carbapenem resistance, future studies should ask in which biological contexts the "
        "prediction is trustworthy, where it is uncertain, and where it is unsafe as a stand-alone "
        "call."
    )

    doc.add_page_break()
    doc.add_heading("Main Figure Legends", level=1)
    add_caption(
        doc,
        "Figure 1. Study design workflow.",
        "Workflow from public-training-only IPM/MEM model development to locked public external "
        "and local clinical validation, mechanism annotation, paired drug-specific MIC error "
        "analysis, clinical warning-rule derivation, and research-use FASTA prototype deployment."
    )
    add_caption(
        doc,
        "Figure 2. Dataset readiness, locked IPM/MEM performance, and genome-defined mechanism evidence.",
        "A, Cohort composition and availability of assemblies and IPM/MEM phenotypes. B, Locked "
        "evaluation performance for IPM and MEM under the same public-only model policy. C, Local "
        "paired IPM/MEM phenotype groups used for within-isolate drug-specific error analysis. "
        "D, Mechanism framework separating OprD-loss, acquired carbapenemase, efflux-regulator, "
        "and AmpC-axis evidence. E, Counts of high-confidence mechanism strata in the local cohort. "
        "F, High-confidence evidence matrix across target mechanism axes. G, PAO1-relative target "
        "gene variation burden, illustrating why broad missense variation was treated as context "
        "rather than as sufficient mechanism-driver evidence. EA, essential agreement; MAE, mean "
        "absolute error; VME, very major error; ME, major error."
    )
    add_caption(
        doc,
        "Figure 3. Mechanism-dependent predictability of IPM and MEM MICs.",
        "A, High-confidence subtype-specific within-one-dilution agreement for IPM and MEM. B, "
        "Subtype-specific bias and error magnitude. C, IPM predicted versus observed MICs by "
        "high-confidence mechanism stratum. D, MEM predicted versus observed MICs by high-confidence "
        "mechanism stratum. E, Subtype-level observed-versus-predicted MIC shift, summarising the "
        "direction and magnitude of endpoint-specific MIC displacement. F, Direction of large "
        "errors by subtype, separating large underprediction, large overprediction, and errors "
        "within two log2 dilutions. OprD-loss and composite backgrounds were interpretable for IPM "
        "but showed systematic MEM underprediction."
    )
    add_caption(
        doc,
        "Figure 4. Paired IPM/MEM errors and sequence-type sensitivity.",
        "A, Paired within-isolate comparison of IPM and MEM absolute MIC error. B, Summary of "
        "MEM-minus-IPM absolute error by mechanism subtype. C, Distribution of local MLST sequence "
        "types in the prediction-analysis subset. D, Leave-one-major-ST-out sensitivity analysis. "
        "E, Exploratory ST-adjusted model comparison showing that mechanism subtype explained more "
        "MEM-minus-IPM error variation than ST group alone. F, ST-level paired error consistency, "
        "showing that positive MEM-minus-IPM absolute error was observed across multiple STs."
    )
    add_caption(
        doc,
        "Figure 5. Clinical interpretation rules and safety-oriented error analysis.",
        "A, High-confidence mechanism-stratum rule table. B, Rule-level performance, including "
        "within-one-dilution agreement and categorical safety errors. C, MEM warning enrichment for "
        "large underprediction among OprD/high-confidence driver backgrounds; standard VME was "
        "broadly high among MEM true non-susceptible isolates and is interpreted as the clinical "
        "safety consequence rather than the primary enrichment signal. D, "
        "Breakpoint-zone analysis distinguishing near-breakpoint drift from far-from-breakpoint "
        "false-susceptible errors. E, Error taxonomy by drug and subtype. F, Safety-gate "
        "proof-of-concept showing VME avoidance after withholding susceptible calls in warning strata."
    )
    add_caption(
        doc,
        "Figure 6. FASTA-based prototype for predictability-aware IPM/MEM reporting.",
        "A, Upload interface for assembly FASTA analysis. B, Example report showing assembly QC, "
        "target-locus k-mer screening, FASTA-derived prototype IPM/MEM MIC estimates, mechanism-"
        "defined predictability stratum, and drug-specific reliability flags. The prototype is "
        "research-use only and is not a validated clinical AST system."
    )

    add_main_tables(doc)
    return doc


def supplementary_doc() -> Document:
    doc = setup_doc()
    add_title(doc, TITLE, f"Supplementary figure legends and supplementary table notes; generated {DATE}")

    doc.add_heading("Supplementary Figure Legends", level=1)
    add_caption(
        doc,
        "Supplementary Figure S1. Unified public-only model policy and public-training apparent performance.",
        "A, Locked IPM/MEM model policy, including breakpoint, model class, feature set, gate "
        "threshold, hard-gate rule, snapping rule, and cap policy. B, Public-training apparent "
        "performance used as a smoke check only; locked validation cohorts were not used for model "
        "training or policy selection."
    )
    add_caption(
        doc,
        "Supplementary Figure S2. First-pass strict mechanism subtype landscape.",
        "A, Counts of strict mechanism subtypes in locked local and locked public external cohorts. "
        "B, Distribution of local IPM/MEM phenotype groups across strict subtypes. C, Strict evidence "
        "matrix showing OprD severe loss, acquired carbapenemase, AmpC-associated genotype, and "
        "efflux-regulator evidence."
    )
    add_caption(
        doc,
        "Supplementary Figure S3. Target-gene variation burden across the mechanism panel.",
        "A, PAO1-relative variant severity landscape across OprD, efflux-regulator, and AmpC-axis "
        "target genes. B, Target-region row composition, distinguishing non-synonymous, regulatory, "
        "disruptive, and context variation. These data support the two-layer mechanism framework."
    )
    add_caption(
        doc,
        "Supplementary Figure S4. Full subtype-specific error distributions.",
        "A, Strict subtype error points by drug. B, High-confidence subtype error points by drug. "
        "The figure complements the main predictability panels by showing isolate-level dispersion "
        "and outliers."
    )
    add_caption(
        doc,
        "Supplementary Figure S5. MLST and sequence-type sensitivity analyses.",
        "A, Local ST counts in the prediction-analysis subset. B, Leave-one-major-ST-out sensitivity "
        "analysis for the paired MEM-minus-IPM error signal. C, Exploratory ST-adjusted model "
        "comparisons for mechanism subtype, ST group, and combined models."
    )
    add_caption(
        doc,
        "Supplementary Figure S6. Calibration and susceptible-call risk.",
        "A, MEM calibration summary by high-confidence mechanism subtype. B, Binned predicted "
        "probability of non-susceptibility versus observed NS rate. C, False-susceptible risk among "
        "MEM predicted susceptible calls, stratified by mechanism subtype and susceptible-call "
        "confidence."
    )
    add_caption(
        doc,
        "Supplementary Figure S7. Safety-gate evaluation, error taxonomy, and false-susceptible cases.",
        "A, Safety-gate scenarios comparing no gate, mechanism-warning susceptible gate, and "
        "mechanism-or-low-margin gate. B, Full error taxonomy summary by drug and high-confidence "
        "subtype. C, False-susceptible case overview with mechanism and prediction context."
    )

    doc.add_heading("Supplementary Table Notes", level=1)
    table_notes = [
        (
            "Supplementary Table S1. Model policy and endpoint definitions.",
            "Locked public-only IPM/MEM model policies, endpoint breakpoints, model classes, feature sets, hard-gate threshold, snapping rule, and public-training apparent summary metrics.",
        ),
        (
            "Supplementary Table S2. Full locked evaluation by cohort and drug.",
            "Complete locked evaluation summaries and isolate-level prediction errors for IPM and MEM in locked public external and locked local validation cohorts.",
        ),
        (
            "Supplementary Table S3. Local paired IPM/MEM phenotype-group metrics.",
            "Local IPM/MEM paired phenotype groups and endpoint-specific performance metrics within each group.",
        ),
        (
            "Supplementary Table S4. First-pass strict mechanism evidence.",
            "Per-isolate strict mechanism evidence and strict subtype assignments, including OprD severe loss, acquired carbapenemase, AmpC-associated genotype, efflux-regulator evidence, and broad context fields.",
        ),
        (
            "Supplementary Table S5. High-confidence mechanism evidence per isolate.",
            "Per-isolate high-confidence mechanism evidence, refined mechanism subtype assignments, and joined IPM/MEM prediction-error context.",
        ),
        (
            "Supplementary Table S6. Target-gene variant summary.",
            "Target-gene and target-row variation summaries across OprD, AmpC-axis, and efflux-regulator loci, supporting the distinction between broad context variation and high-confidence driver evidence.",
        ),
        (
            "Supplementary Table S7. High-confidence paired IPM/MEM statistical tests.",
            "Paired Wilcoxon, binomial, global, and subtype-level statistical summaries supporting drug-specific IPM/MEM predictability differences.",
        ),
        (
            "Supplementary Table S8. MLST and ST sensitivity.",
            "Exact PubMLST calls, ST counts, paired error summaries by ST, leave-one-ST-out analyses, and exploratory ST-adjusted models.",
        ),
        (
            "Supplementary Table S9. Clinical warning rule full table.",
            "Full high-confidence clinical warning-rule table, rule-level performance summaries, and MEM warning enrichment tests.",
        ),
        (
            "Supplementary Table S10. False-susceptible cases.",
            "Case-level false-susceptible predictions and top error cases, including actual MIC, predicted MIC, mechanism subtype, rule label, and target-gene severity context.",
        ),
        (
            "Supplementary Table S11. Calibration and breakpoint-distance summaries.",
            "Calibration summaries, probability bins, susceptible-call risk, breakpoint-zone error summaries, and phenotype-combination error summaries.",
        ),
        (
            "Supplementary Table S12. Error taxonomy and safety-gate evaluation.",
            "Full error taxonomy rows and summaries plus safety-gate evaluation outputs for no-gate and mechanism-aware withholding scenarios.",
        ),
    ]
    for label, text in table_notes:
        add_caption(doc, label, text)

    doc.add_heading("Supplementary Table Files", level=1)
    add_supplementary_table_file_index(doc)

    doc.add_heading("Supplementary Files Already Prepared", level=1)
    add_bullet(doc, "Main figure editable PDFs: 投稿/main_figures_editable_pdf/")
    add_bullet(doc, "Supplementary figure editable PDFs: 投稿/supplementary_figures_editable_pdf/")
    add_bullet(doc, "Main and supplementary manuscript tables: 投稿/manuscript_tables/")
    add_bullet(doc, "All tables combined workbook: 投稿/manuscript_tables/IPM-GPT_all_manuscript_tables_2026-06-07.xlsx")
    add_bullet(doc, "Research-use web prototype upload package: 投稿/huggingface_docker_upload/")

    doc.add_heading("Abbreviations", level=1)
    add_p(
        doc,
        "AST, antimicrobial susceptibility testing; EA, essential agreement; IPM, imipenem; MAE, "
        "mean absolute error; ME, major error; MEM, meropenem; MIC, minimum inhibitory concentration; "
        "MLST, multilocus sequence typing; NS, non-susceptible; S, susceptible; ST, sequence type; "
        "VME, very major error; WGS, whole-genome sequencing."
    )
    return doc


def main() -> None:
    OUT.mkdir(parents=True, exist_ok=True)
    main_doc = main_manuscript()
    supp_doc = supplementary_doc()
    main_path = OUT / f"carbapenem_predictability_full_manuscript_draft_{DATE}.docx"
    supp_path = OUT / f"carbapenem_predictability_supplementary_legends_tables_{DATE}.docx"
    main_doc.save(main_path)
    supp_doc.save(supp_path)
    print(main_path)
    print(supp_path)


if __name__ == "__main__":
    main()
